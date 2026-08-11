#!/usr/bin/env python3
"""ประกอบบล็อก markdown เป็นไฟล์ Word .docx ด้วย python-docx"""

from __future__ import annotations

import sys
from pathlib import Path

from mdblocks import fit_inches, is_caption, iter_inline, split_table_row

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.image.image import Image as DocxImage
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:  # pragma: no cover
    sys.exit("ต้องติดตั้ง python-docx ก่อน:  pip install python-docx")

INK = RGBColor(0x1A, 0x1A, 0x1A)
HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F)
CAPTION_COLOR = RGBColor(0x44, 0x44, 0x44)
HEADER_FILL = "E8EEF7"

# ลำดับ element ที่อยู่ถัดจาก w:szCs ตาม schema ของ w:rPr
SZCS_SUCCESSORS = (
    "w:highlight", "w:u", "w:effect", "w:bdr", "w:shd", "w:fitText",
    "w:vertAlign", "w:rtl", "w:cs", "w:em", "w:lang", "w:eastAsianLayout",
    "w:specVanish", "w:oMath",
)


# ---------------------------------------------------------------- ฟอนต์ไทย

def _set_cs_font(rPr, font: str, size_pt: float | None, bold: bool | None) -> None:
    """ตั้งฟอนต์และขนาดของ complex script (w:cs / w:szCs / w:bCs)

    ถ้าไม่ตั้ง Word จะใช้ค่า default กับสระและวรรณยุกต์ไทย ทำให้ตัวอักษรลอยผิดตำแหน่ง
    """
    rFonts = rPr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), font)
    if size_pt is not None:
        szCs = rPr.find(qn("w:szCs"))
        if szCs is None:
            szCs = OxmlElement("w:szCs")
            rPr.insert_element_before(szCs, *SZCS_SUCCESSORS)
        szCs.set(qn("w:val"), str(int(round(size_pt * 2))))
    if bold is not None:
        rPr.get_or_add_bCs().set(qn("w:val"), "1" if bold else "0")


def style_run(run, font: str, size_pt: float, *, bold=None, italic=None, color=None, mono=None):
    face = mono or font
    run.font.name = face
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    _set_cs_font(run._element.get_or_add_rPr(), face, size_pt, bold)
    return run


def style_style(style, font: str, size_pt: float, *, bold=None, color=None):
    style.font.name = font
    style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    _set_cs_font(style.element.get_or_add_rPr(), font, size_pt, bold)


# ---------------------------------------------------------------- ตัวเขียน

class DocxWriter:
    def __init__(self, args):
        self.font = args.font
        self.mono = args.mono_font
        self.size = args.size
        # Word เกลี่ยช่องไฟระหว่างอักษรไทยได้ จัดชิดขอบสองด้านจึงสวยตามธรรมเนียมเอกสารไทย
        self.align = (WD_ALIGN_PARAGRAPH.LEFT if args.align == "left"
                      else WD_ALIGN_PARAGRAPH.JUSTIFY)
        self.max_w = args.max_width
        self.max_h = args.max_height
        self.scale = args.scale
        self.doc = Document()
        self._setup_page(args)
        self._setup_styles()

    # --- โครงหน้ากระดาษและสไตล์ ---

    def _setup_page(self, args):
        section = self.doc.sections[0]
        section.page_width = Inches(8.27)   # A4
        section.page_height = Inches(11.69)
        for side in ("top", "bottom", "left", "right"):
            setattr(section, f"{side}_margin", Inches(args.margin))
        self.max_w = min(self.max_w, 8.27 - 2 * args.margin)
        self._page_number_footer(section)

    def _page_number_footer(self, section):
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = style_run(p.add_run(), self.font, self.size - 2, color=CAPTION_COLOR)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        for el in (begin, instr, end):
            run._r.append(el)

    def _setup_styles(self):
        normal = self.doc.styles["Normal"]
        style_style(normal, self.font, self.size, color=INK)
        pf = normal.paragraph_format
        pf.space_after = Pt(8)
        pf.line_spacing = 1.15

        for level, (pt, space_before) in {
            1: (self.size + 8, 0), 2: (self.size + 4, 16),
            3: (self.size + 2, 12), 4: (self.size + 1, 10),
        }.items():
            st = self.doc.styles[f"Heading {level}"]
            style_style(st, self.font, pt, bold=True, color=HEADING_COLOR)
            st.paragraph_format.space_before = Pt(space_before)
            st.paragraph_format.space_after = Pt(6)
            st.paragraph_format.keep_with_next = True

    # --- บล็อกแต่ละชนิด ---

    def add_runs(self, paragraph, text: str, size: float, *, base_bold=False, color=None):
        for chunk, fmt in iter_inline(text):
            if not chunk:
                continue
            style_run(
                paragraph.add_run(chunk.replace("<br/>", "\n").replace("<br>", "\n")),
                self.font, size,
                bold=fmt.get("bold", base_bold or None),
                italic=fmt.get("italic"),
                color=color,
                mono=self.mono if fmt.get("code") else None,
            )

    def heading(self, level: int, text: str):
        p = self.doc.add_heading(level=min(level, 4))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
        size = {1: self.size + 8, 2: self.size + 4, 3: self.size + 2}.get(level, self.size + 1)
        self.add_runs(p, text, size, base_bold=True, color=HEADING_COLOR)
        if level == 1:
            p.paragraph_format.space_after = Pt(14)

    def paragraph(self, text: str):
        caption = is_caption(text)
        p = self.doc.add_paragraph()
        p.alignment = self.align
        self.add_runs(p, text, self.size - 1 if caption else self.size,
                      color=CAPTION_COLOR if caption else None)
        if caption:
            pf = p.paragraph_format
            pf.left_indent = Inches(0.3)
            pf.right_indent = Inches(0.3)
            pf.space_before = Pt(2)
            pf.space_after = Pt(14)

    def bullets(self, items: list[dict]):
        for item in items:
            style = "List Number" if item["ordered"] else "List Bullet"
            p = self.doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Inches(0.3 + 0.3 * item["indent"])
            p.alignment = self.align
            self.add_runs(p, item["text"], self.size)

    def code(self, payload: dict):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        style_run(p.add_run(payload["text"]), self.font, self.size - 2,
                  color=CAPTION_COLOR, mono=self.mono)
        self._shade(p._p.get_or_add_pPr(), "F3F4F6")

    def rule(self):
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "C9CED6")):
            bottom.set(qn(k), v)
        borders.append(bottom)
        pPr.append(borders)

    def image(self, path: Path):
        if not path.is_file():
            self.paragraph(f"*[ไม่พบไฟล์ภาพ: {path.name}]*")
            return
        meta = DocxImage.from_file(str(path))
        w_in, _ = fit_inches(meta.px_width, meta.px_height, self.scale, self.max_w, self.max_h)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(path), width=Inches(w_in))

    def table(self, rows: list[str]):
        header = split_table_row(rows[0])
        body = [split_table_row(r) for r in rows[2:] if r.strip("| ")]
        cols = max([len(header)] + [len(r) for r in body])

        table = self.doc.add_table(rows=1 + len(body), cols=cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        self._table_width_pct(table)

        for idx, text in enumerate(header + [""] * (cols - len(header))):
            cell = table.rows[0].cells[idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.add_runs(p, text, self.size - 1, base_bold=True)
            self._shade(cell._tc.get_or_add_tcPr(), HEADER_FILL)
        self._repeat_header(table.rows[0])

        for r, cells in enumerate(body, start=1):
            for c in range(cols):
                self.add_runs(table.rows[r].cells[c].paragraphs[0],
                              cells[c] if c < len(cells) else "",
                              self.size - 1)

        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)

        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # --- helper XML ---

    @staticmethod
    def _shade(parent, fill: str):
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        parent.append(shd)

    @staticmethod
    def _repeat_header(row):
        trPr = row._tr.get_or_add_trPr()
        el = OxmlElement("w:tblHeader")
        el.set(qn("w:val"), "true")
        trPr.append(el)

    @staticmethod
    def _table_width_pct(table):
        tblPr = table._tbl.tblPr
        for tag, attrs in (("w:tblW", {"w:w": "5000", "w:type": "pct"}),
                           ("w:tblLayout", {"w:type": "autofit"})):
            el = tblPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                tblPr.append(el)
            for k, v in attrs.items():
                el.set(qn(k), v)

    # --- ประกอบทั้งเอกสาร ---

    def build(self, blocks, resolve_image):
        for kind, payload in blocks:
            if kind == "heading":
                self.heading(payload["level"], payload["text"])
            elif kind == "para":
                self.paragraph(payload)
            elif kind == "list":
                self.bullets(payload)
            elif kind == "table":
                self.table(payload)
            elif kind == "image":
                self.image(resolve_image(payload["src"]))
            elif kind == "code":
                self.code(payload)
            elif kind == "hr":
                self.rule()

    def save(self, path: Path, title: str, source: Path):
        props = self.doc.core_properties
        props.title = title
        props.comments = f"สร้างจาก {source.as_posix()} โดยสกิล srs-export-overview"
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))


def write_docx(blocks, args, out: Path, title: str, source: Path, resolve_image) -> Path:
    writer = DocxWriter(args)
    writer.build(blocks, resolve_image)
    writer.save(out, title, source)
    return out
